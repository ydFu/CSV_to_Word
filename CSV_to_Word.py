import csv
from docx import Document

# read the csv file
with open('data.csv', 'r', encoding='utf-8') as csv_file:
    csv_reader = csv.DictReader(csv_file)
    data = [row for row in csv_reader]

# read the template
doc = Document('input.docx')

# replace the placeholders with the data from the csv file
for row in data:
    for p in doc.paragraphs:
        if '&&1' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '&&1' in inline[i].text:
                    text = inline[i].text.replace('&&1', row['&&1'])
                    inline[i].text = text
        if '&&2' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if '&&2' in inline[i].text:
                    text = inline[i].text.replace('&&2', row['&&2'])
                    inline[i].text = text
        if '&&3' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if '&&3' in inline[i].text:
                    text = inline[i].text.replace('&&3', row['&&3'])
                    inline[i].text = text
        if '&&4' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if '&&4' in inline[i].text:
                    text = inline[i].text.replace('&&4', row['&&4'])
                    inline[i].text = text
        if '&&5' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if '&&5' in inline[i].text:
                    text = inline[i].text.replace('&&5', row['&&5'])
                    inline[i].text = text
        if '&&6' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if '&&6' in inline[i].text:
                    text = inline[i].text.replace('&&6', row['&&6'])
                    inline[i].text = text
        if '&&7' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if '&&7' in inline[i].text:
                    text = inline[i].text.replace('&&7', row['&&7'])
                    inline[i].text = text
        if '&&8' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if '&&8' in inline[i].text:
                    text = inline[i].text.replace('&&8', row['&&8'])
                    inline[i].text = text
        if '&&9' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if '&&9' in inline[i].text:
                    text = inline[i].text.replace('&&9', row['&&9'])
                    inline[i].text = text

    for t in doc.tables:
        for r in t.rows:
            for cell in r.cells:
                if '&&1' in cell.text:
                    cell.text = cell.text.replace('&&1', row['&&1'])
                if '&&2' in cell.text:
                    cell.text = cell.text.replace('&&2', row['&&2'])
                if '&&3' in cell.text:
                    cell.text = cell.text.replace('&&3', row['&&3'])
                if '&&4' in cell.text:
                    cell.text = cell.text.replace('&&4', row['&&4'])
                if '&&5' in cell.text:
                    cell.text = cell.text.replace('&&5', row['&&5'])
                if '&&6' in cell.text:
                    cell.text = cell.text.replace('&&6', row['&&6'])
                if '&&7' in cell.text:
                    cell.text = cell.text.replace('&&7', row['&&7'])
                if '&&8' in cell.text:
                    cell.text = cell.text.replace('&&8', row['&&8'])
                if '&&9' in cell.text:
                    cell.text = cell.text.replace('&&9', row['&&9'])

# save the file
doc.save('output.docx')
